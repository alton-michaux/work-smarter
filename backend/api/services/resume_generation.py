import io
import json
import logging

import openai
from django.conf import settings
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

IMPROVE_PROMPT = """\
You are a professional resume writer. Rewrite the provided resume to be significantly stronger.

Use the work history context to:
- Add relevant accomplishments from task history that are absent or understated
- Rewrite weak bullets using action verbs and quantifiable impact
- Improve clarity, flow, and professional tone throughout
- Apply the analysis suggestions if provided

Return ONLY the improved resume in markdown. Use this exact structure:
# Full Name
## Contact
Email | Phone | Location | LinkedIn

## Summary
2-3 sentence professional summary.

## Experience
### Job Title — Company (Month Year – Month Year or Present)
- Bullet with action verb and impact
- Another bullet

## Education
### Degree — School (Year)

## Skills
Skill 1, Skill 2, Skill 3

Rules:
- Preserve all factual details (names, dates, companies) exactly as given.
- Do not invent companies, roles, dates, or credentials.
- Return only the markdown resume. No commentary, no preamble.\
"""

SCRATCH_PROMPT = """\
You are a professional resume writer. Build a professional resume from the work history below.

The user has not provided an existing resume. Infer experience from their completed task history and projects.

If a USER PROFILE section is provided, use that exact information (name, email, phone, location) in the contact header instead of placeholders.

Return ONLY a resume in markdown. Use this exact structure:
# Full Name
## Contact
Email | Phone | Location | LinkedIn

## Professional Summary
2-3 sentence summary inferred from their work.

## Experience
### Inferred Job Title — Project/Company Name (Date Range)
- Transform completed tasks into impactful bullets starting with action verbs

## Skills
Infer from task descriptions and project types.

## Education
### Degree — School (Year)
(Leave as placeholder if not determinable)

Rules:
- Group tasks logically into jobs or projects. If tasks have no projects, only infer the skills gained from those tasks.
- Use the "Date Range" field from each project section exactly as provided — do not change or fabricate dates.
- Use action verbs: Built, Designed, Led, Implemented, Improved, etc.
- Do not fabricate specific metrics unless clearly inferable.
- Return only the markdown resume. No commentary.\
"""


def _call_groq(system_prompt: str, user_content: str) -> str:
    openai.api_key = settings.GROQ_API_KEY
    openai.api_base = "https://api.groq.com/openai/v1"

    response = openai.ChatCompletion.create(
        model="llama-3.3-70b-versatile",
        max_tokens=3000,
        temperature=0.4,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
    )
    return response.choices[0].message.content.strip()


def generate_improved_resume(resume_text: str, task_context: str, analysis: dict | None = None) -> str:
    parts = [f"=== EXISTING RESUME ===\n{resume_text}", f"=== WORK HISTORY ===\n{task_context}"]
    if analysis:
        hints = {
            "suggestions": analysis.get("suggestions", []),
            "accomplishments": analysis.get("accomplishments", []),
            "bullet_rewrites": analysis.get("bullet_rewrites", {}),
        }
        parts.append(f"=== ANALYSIS SUGGESTIONS ===\n{json.dumps(hints, indent=2)}")
    user_content = "\n\n".join(parts)
    return _call_groq(IMPROVE_PROMPT, user_content)


def generate_resume_from_scratch(task_context: str, user_info: dict | None = None) -> str:
    parts = []
    if user_info:
        info_lines = []
        if user_info.get('name'):
            info_lines.append(f"Name: {user_info['name']}")
        if user_info.get('email'):
            info_lines.append(f"Email: {user_info['email']}")
        if user_info.get('phone'):
            info_lines.append(f"Phone: {user_info['phone']}")
        if user_info.get('location'):
            info_lines.append(f"Location: {user_info['location']}")
        if info_lines:
            parts.append("=== USER PROFILE ===\n" + '\n'.join(info_lines))
    parts.append(f"=== WORK HISTORY ===\n{task_context}")
    user_content = "\n\n".join(parts)
    return _call_groq(SCRATCH_PROMPT, user_content)


def content_to_docx(markdown: str) -> bytes:
    """Convert a simple markdown resume to a DOCX byte stream."""
    doc = Document()

    # Tighten default margins
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Pt(36)
    section.left_margin = section.right_margin = Pt(54)

    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            # Inline **bold** stripping for plain-text fallback
            text = stripped.replace("**", "")
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
